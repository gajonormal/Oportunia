from docx import Document
from docx.shared import Pt
import sys

def add_final_delivery(file_path):
    doc = Document(file_path)
    
    # Adicionar o Heading Principal
    doc.add_heading('7. Funcionalidades da Última Entrega (Entrega 3)', level=1)
    
    # Parágrafo introdutório
    doc.add_paragraph(
        "Nesta última iteração do projeto Oportunia, o foco principal recaiu sobre a estabilização da arquitetura em Cloud (Azure), "
        "o enriquecimento da integração com a Inteligência Artificial, e o tratamento de limitações técnicas inesperadas encontradas ao longo da investigação. "
        "Abaixo detalham-se as funcionalidades concluídas, as barreiras encontradas e a forma como o sistema foi readaptado com sucesso."
    )
    
    # 7.1
    doc.add_heading('7.1. Recuperação Segura de Password', level=2)
    doc.add_paragraph(
        "Foi implementado um fluxo completo para a recuperação de credenciais esquecidas. O sistema utiliza a biblioteca 'itsdangerous' "
        "do Python para gerar tokens únicos e com limite temporal, enviados de forma segura para a caixa de correio do utilizador via 'smtplib'. "
        "Na camada frontend (HTML/JS), adicionou-se lógica de interceção de cliques múltiplos para mitigar problemas de spam caso o botão fosse acionado consecutivamente."
    )
    
    # 7.2
    doc.add_heading('7.2. Extração Inteligente de CV: O Desvio das APIs Estatais', level=2)
    doc.add_paragraph(
        "A arquitetura original (conforme descrita nas Fases 1 e 2 do relatório) previa a integração sistemática com APIs do programa Erasmus+ e do IEFP/FCT. "
        "O que não conseguimos fazer como pretendíamos: após pesquisa extensiva, constatou-se que grande parte destas APIs europeias e estatais são de acesso restrito (requerendo parcerias institucionais morosas) "
        "ou possuem documentação obsoleta e sem suporte a chaves públicas de teste para fins académicos."
    )
    doc.add_paragraph(
        "A Solução Encontrada (Pivot): Para contornar esta restrição de obtenção massiva de dados do exterior, a estratégia do projeto pivotou para o enriquecimento da informação vinda do utilizador. "
        "Desenvolveu-se a funcionalidade de 'Extração Inteligente de Currículos'. O estudante arrasta um ficheiro PDF para a plataforma; o backend (usando a biblioteca 'PyPDF2') "
        "descodifica o texto em bruto, e injeta-o numa diretriz rigorosa ao 'GPT-4o'. A IA analisa o documento não-estruturado e devolve, no formato JSON estrito, "
        "as 'Competências Técnicas', 'Idiomas' e 'Projetos' relevantes, populando os inputs do ecrã de perfil dinamicamente em segundos. Isto resultou numa experiência de utilizador muito superior."
    )
    
    # 7.3
    doc.add_heading('7.3. Integração com Azure Blob Storage e Gestão RGPD', level=2)
    doc.add_paragraph(
        "Para garantir o armazenamento seguro dos currículos gerados ou carregados pelos alunos, integrou-se o 'Azure Storage Blob'. "
        "A comunicação foi feita através da biblioteca oficial 'azure-storage-blob'. Quando o PDF é processado, não se perde na memória RAM: "
        "é guardado num contentor seguro na Cloud, e o seu localizador (nome do blob) é anexado à conta do aluno no 'CosmosDB'."
    )
    doc.add_paragraph(
        "Em conformidade com as boas práticas de privacidade e do RGPD, a plataforma permite não só a consulta / download do CV através de uma rota dedicada ('GET /api/cv/download'), "
        "como também a total eliminação física do ficheiro armazenado na Cloud e do registo na BD ('DELETE /api/cv/remove') mediante confirmação."
    )
    
    # 7.4
    doc.add_heading('7.4. Otimização do Motor IA e do Contentor Docker em Produção', level=2)
    doc.add_paragraph(
        "Verificou-se que enviar dezenas de vagas indiscriminadamente para o GPT-4o analisar e justificar criava gargalos de performance drásticos (e estourava limites de tokens API). "
        "Implementou-se assim um algoritmo 'Pré-Rank' em Python que faz uma triagem preliminar através das Tags e Localizações do aluno. "
        "Apenas as vagas mais promissoras (Top 20) são passadas à IA para avaliação profunda de match."
    )
    doc.add_paragraph(
        "Adicionalmente, verificámos falhas no ambiente de Produção (Azure App Service) quando a IA demorava a pensar. O servidor nativo Gunicorn encerrava o processo ao fim de 30 segundos (default). "
        "Para suportar a assincronicidade das respostas longas do GPT, reescreveu-se o comando de iniciação no 'Dockerfile' para acomodar um timeout de 120s ('--timeout 120'), o que resolveu definitivamente as quebras de conexão."
    )
    
    doc.save(file_path)
    print("Sucesso! Relatorio atualizado.")

if __name__ == "__main__":
    add_final_delivery(sys.argv[1])
