from docx import Document
import sys

def delete_paragraph(paragraph):
    p = paragraph._element
    if p.getparent() is not None:
        p.getparent().remove(p)
    p._p = p._element = None

def build_final_report(main_file, source_file):
    doc_main = Document(main_file)
    doc_source = Document(source_file)
    
    # 1. Eliminar do doc principal os paragrafos do capitulo 7 (criados antes) e Referencias
    found_chapter_7 = False
    for p in list(doc_main.paragraphs):
        if p.text.startswith("7. Funcionalidades da Última Entrega"):
            found_chapter_7 = True
        
        if p.text.strip().lower() == "referências" or p.text.startswith("Referências"):
            found_chapter_7 = True
            
        if found_chapter_7:
            delete_paragraph(p)

    # 2. Escrever a secção final (nova)
    doc_main.add_heading('7. Integrações e Funcionalidades da Última Entrega (Entrega 3)', level=1)
    
    doc_main.add_paragraph(
        "Nesta terceira e última fase do projeto, o foco transitou do protótipo de interface para um ecossistema cloud totalmente integrado. "
        "A aplicação passou a ter o seu frontend, que anteriormente consistia apenas num layout estático de perfis, "
        "totalmente acoplado à lógica backend (Flask) e à base de dados NoSQL (CosmosDB). Isto garantiu que ações como alteração de tags "
        "e definições no ecrã de Perfil passem a modificar instantaneamente e de forma real o documento JSON do utilizador na Cloud, "
        "com impacto direto na Dashboard de vagas recomendadas."
    )
    
    doc_main.add_heading('7.1. Extração Inteligente de CV: O Desvio Estratégico das APIs', level=2)
    doc_main.add_paragraph(
        "A arquitetura original previa a integração constante com APIs institucionais de emprego e estágios (IEFP, Erasmus+, FCT). "
        "Contudo, a equipa defrontou-se com uma limitação grave: grande parte destas APIs são de acesso restrito (requerendo chaves B2B), "
        "possuem documentação fragmentada ou não disponibilizam conectores Sandbox para testes letivos."
    )
    doc_main.add_paragraph(
        "Como solução, a equipa pivotou a estratégia de enriquecimento de dados da Fonte para o Utilizador. Desenvolveu-se a funcionalidade de 'Extração Inteligente de Currículos': "
        "O aluno submete o seu PDF na interface. O backend decodifica o texto em bruto com a biblioteca PyPDF2 e encarrega o motor GPT-4o de estruturar a "
        "carreira do aluno em formato JSON estrito, injetando automaticamente as 'Competências', 'Idiomas' e 'Projetos' na vista dinâmica."
    )
    
    doc_main.add_heading('7.2. Azure Blob Storage e Gestão de Privacidade (RGPD)', level=2)
    doc_main.add_paragraph(
        "De forma a suportar uploads sem perder os ficheiros na efemeridade dos contentores, "
        "integrou-se o Azure Blob Storage. Os PDFs são guardados de forma segura num contentor na nuvem, com apenas a referência string a "
        "ser guardada no perfil do aluno no CosmosDB. Foi também incluída na interface a funcionalidade de Descarregar o CV e de "
        "Apagar o Ficheiro, chamando o método de eliminação física do SDK da Microsoft, garantindo total conformidade com o RGPD."
    )

    # 3. Copiar os conteudos detalhados do relatorio complementar
    copy_mode = False
    for p in doc_source.paragraphs:
        if p.text.startswith("6.9.1. Recuperação"):
            copy_mode = True
            
        if copy_mode:
            text = p.text
            if text.startswith("6.9.1"): text = text.replace("6.9.1", "7.3")
            elif text.startswith("6.9.2"): text = text.replace("6.9.2", "7.4")
            elif text.startswith("6.9.3"): text = text.replace("6.9.3", "7.5")
            elif text.startswith("6.9.4"): text = text.replace("6.9.4.", "7.6.")
            elif text.startswith("7."): text = text.replace("7.", "8.")
            elif text.startswith("8."): text = text.replace("8.", "9.")
            elif text.startswith("9."): text = text.replace("9.", "10.")
                
            if p.style.name.startswith('Heading 1'):
                doc_main.add_heading(text, level=1)
            elif p.style.name.startswith('Heading 2'):
                doc_main.add_heading(text, level=2)
            elif p.style.name.startswith('Heading'):
                doc_main.add_heading(text, level=3)
            else:
                if text.strip() != "":
                    doc_main.add_paragraph(text)
                    
    # 4. Voltar a adicionar as Referencias que apaguei no passo 1
    doc_main.add_heading('Referências', level=1)
    doc_main.add_paragraph("Documentação Oficial Microsoft Azure (App Service, Blob Storage, CosmosDB, Functions).")
    doc_main.add_paragraph("Documentação Oficial OpenAI API (GPT-4o, Tokens).")
    doc_main.add_paragraph("Especificações Flask e Documentação Jinja2.")

    doc_main.save(main_file)
    print("Relatório completo montado com sucesso.")

if __name__ == "__main__":
    build_final_report(sys.argv[1], sys.argv[2])
